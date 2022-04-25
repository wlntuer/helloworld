import re
import jieba
import docx
from wordcloud import WordCloud
import matplotlib.pyplot as plt


#输入参数path为word文档所在路径
def readFile(path):
    file=docx.Document(path)
    # 读取每一段的内容
    article = ' '
    for para in file.paragraphs:
        if para != '/n':
            article = article + '' + para.text
    # file.paragraphs只能读取到文档中正常段落中的文本信息，却无法读取到文档中表格里的内容，所以我们还需要继续将表格中的内容也读取出来进行合并
    # 获取文档中表格信息
    tables = file.tables  # 获取文档中所有表格对象的列表
    for i in range(len(tables)):
        table0 = tables[i]  # 获取表格对象
        # 获取一个表格的所有单元格
        cells = table0._cells
        # 获取单元格内所有文字信息
        cells_string = [cell.text.replace('\n', ' ').replace('\r', ' ') for cell in cells]
        # print(cells_string)
        for i in range(len(cells_string)):
            article = article + cells_string[i].strip()
    return article


def wordFrequency(a):
    # print(article)
    # print(len(article))
    # 全模式 cut_all=True
    # seq_list=jieba.cut(article,cut_all=True)
    # print(seq_list) #<generator object Tokenizer.cut at 0x0000026EB6F0CD58>
    # print('全模式',list(seq_list))
    # 精确模式 （默认模式） cut_all =False
    # seq_list=jieba.cut(article,cut_all=False)
    # print('精确模式',list(seq_list))
    # 搜索引擎模式 cut_for_search
    seq_list = jieba.cut_for_search(article, )
    # print('搜索引擎模式',list(seq_list))
    #通过符号将字符串分割开
    articleSplit = re.split(" |,|\.|，|。|;|；|、|：|:|：", article)
    sep_list = " ".join(articleSplit)  # 转为字符串
    # print(sep_list)
    articleSplit = seq_list
    d = {}
    #需要忽略的符号集合
    fuhao = [',', '.', ' ', ':', '：', '，', '。', '[', ']', ';', '；', '(', ')', '（', '）', '!', '！']
    #统计词频
    for key in articleSplit:
        if key in d.keys():
            d[key] = d[key] + 1
        else:
            if (len(key.strip()) == 0 or key in fuhao):
                continue
            d[key] = 1
    wordlist = ''
    for key in d:
        # print(key,d[key])
        wordlist = wordlist + ' ' + key
    # print(wordlist)
    #根据词频进行排序
    sort_words = sorted(d.items(), key=lambda x: x[1], reverse=True)
    print(sort_words)
    return wordlist


def wordcloudGenerate(wordlist):
    wc = WordCloud(background_color='white',max_words=2000,width=800,height=600,relative_scaling=1,
                       max_font_size=80,random_state=40,font_path="simhei.ttf")
    # 生成词云
    wc.generate(wordlist)
    # 在只设置mask的情况下，你将会得到一个拥有图片形状的词云
    plt.rcParams['savefig.dpi'] = 800  # 图片像素
    plt.rcParams['figure.dpi'] = 800  # 分辨率
    plt.imshow(wc, interpolation='bilinear')
    plt.axis("off")
    plt.show()


#获取文档对象
path ="C:\\Users\\wlntu\\Desktop\\工作总结\\20221223修订\\2022年度工作总结-财务会计部 20221223 v0.2.docx"

article = readFile(path)
wordlist = wordFrequency(article)
wordcloudGenerate(wordlist)
